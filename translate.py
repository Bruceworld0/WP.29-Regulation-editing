from googletrans import Translator
from openpyxl import Workbook
from openpyxl import load_workbook
from docx import Document

translator = Translator()
'''
write_wb = Workbook()
write_ws = write_wb.active
doc = Document("R048r12e.docx")
paragraphs = [paragraph.text for paragraph in doc.paragraphs]
for i, para in enumerate(paragraphs):
    write_ws.cell(i+1, 1, para)
    try:
        output = translator.translate(para, src='en', dest='ko')
        # print(output.text)
        write_ws.cell(i+1, 2, output.text)
    except :
        pass
'''
load_wb = load_workbook("R048r12e.xlsx", data_only=True)

def store(a, b):
    load_ws = load_wb[a]

    doc = Document(b)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    for i, para in enumerate(paragraphs):
        load_ws.cell(i+1, 1, para)
        
        try:
            output = translator.translate(para, src='en', dest='ko')
            # print(output.text)
            load_ws.cell(i+1, 2, output.text)
        except :
            pass

sheets = ["0601", "0602", "0603", "0604", "0605", "0606", "0607", "0608", "0609", "0610", "0611", "0612", "0613", "0614", "0615", \
          "0701", "0702", "0703", "0704", "0705", "0706", \
          "0801", "0802", "0803"]

documents = ["R048r12am1e.docx", "R048r12am2e.docx", "R048r12am3e.docx", "R048r12am4e.docx", "R048r12am5e.docx", "R048r12am6e.docx", "R048r12am7e.docx", "R048r12am8e.docx", "R048r12am9e.docx", "R048r12am10e.docx", "R048r12am11e.docx", "R048r12am12e.docx", "R048r12am13e.docx", "R048r12am14e.docx", "R048r12am15e.docx",\
             "R048r13am1e.docx", "R048r13am2e.docx", "R048r13am3e.docx", "R048r13am4e.docx", "R048r13am5e.docx", "R048r13am6e.docx",  \
             "R048r14am1e.docx", "R048r14am2e.docx", "R048r14am3e.docx"]

for a, b in zip(sheets, documents):
    store(a,b)

load_wb.save("UNR48_06_07_08.xlsx")

